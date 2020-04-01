from flask import send_from_directory
from docx import Document, styles
from docx.shared import Pt
import os
from datetime import datetime
from app.gdocs_reader import read

def spravkaFieldPicker(date):
    month = date.month
    years = []
    year = date.year
    months = []
    for i in range(6):
        years.append(str(year))
        months.append(str(month))
        month -= 1
        if month < 1:
            month = 12
            year = year - 1
    return [years, months]


def replace_request(words, replacements, doc):
    for j in range(0, len(words)):
        if not replacements[j]:
            replacements[j] = ' '
        for k in range(0, len(doc.tables)):
            for i in doc.tables[k]._cells:
                if words[j] in i.text:
                    i.text = i.text.replace(str(words[j]), str(replacements[j]))

    return doc


def replace_doc(words, replacements, doc):
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = "Times New Roman"
    for j in range(0, len(words)):
        if not replacements[j]:
            replacements[j] = ' '

        for i in doc.paragraphs:
            if words[j] in i.text:
                i.style = style
                i.text = i.text.replace(str(words[j]), str(replacements[j]))
        try:
            if words[j] in doc.tables[0].cell(0, 1).text:
                doc.tables[0].cell(0, 1).text = doc.tables[0].cell(0, 1).text.replace(str(words[j]), str(replacements[j]))
        except Exception:
            continue
    return doc


def five(userid):
    info = read(mode='data')[userid]
    print(info[48])
    doc = Document(os.path.dirname(__file__) + '/files/Form_З_502_1.docx')
    doc = replace_request(words=['{{Сумма_кредита}}', '{{Срок_кредита}}', '{{Первоначальный_взнос}}',
                                 '{{Оценочная стоимость}}',
                                 '{{Сумма_мат_капитала}}', '{{ФИО}}', '{{ФИО_до}}', '{{Серия_паспорта}}',
                                 '{{Номер_паспорта}}', '{{Код_подразделения}}', '{{Кем_выдан}}', '{{Область_проп}}',
                                 '{{Район_проп}}', '{{Город_проп}}', '{{Улица_проп}}', '{{Дом_проп}}',
                                 '{{Кватрира_проп}}', '{{Подъезд_проп}}', '{{Область}}', '{{Район}}', '{{Город}}',
                                 '{{Улица}}', '{{Дом}}', '{{Квартира}}', '{{Телефон}}', '{{email}}',
                                 '{{Снилс}}', '{{Дети}}', '{{Кому_меньше}}', '{{Наименование_работодателя}}',
                                 '{{Должность}}', '{{Месяц_и_год_устройства}}', '{{Сайт_работодателя}}', '{{ИНН}}',
                                 '{{Дата_выдачи}}'],
                          replacements=[info[53], info[55], info[54], info[42].split(', ')[1], '', info[3], info[20],
                                        info[5].split(' ')[0], info[5].split(' ')[1], '', info[7],
                                        info[9].split(' обл.')[0], '', info[9].split('г. ')[1].split(', ул.')[0],
                                        info[9].split('ул. ')[1].split(', д.')[0],
                                        info[9].split(', д.')[1].split(', кв.')[0], info[9].split('кв. ')[1], '',
                                        info[10].split(' обл.')[0], '', info[10].split('г. ')[1].split(', ул.')[0],
                                        info[10].split('ул. ')[1].split(', д.')[0],
                                        info[10].split(', д.')[1].split(', кв.')[0], info[10].split('кв. ')[1],
                                        info[11], info[12], info[8], len(info[23].split(', ')),
                                        len(info[23].split(', ')), info[13], info[25],
                                        info[71].split('.')[1] + ', ' + info[71].split('.')[2], info[33],
                                        info[14], info[6]], doc=doc)
    doc.save(os.path.dirname(__file__) + '/files/downloaded.docx')
    return send_from_directory(directory=os.path.abspath(os.path.dirname(__file__) + '/files'),
                                   filename='downloaded.docx')


def four(userid):
    date = datetime.today()
    info = read(mode='data')[userid]
    doc = Document(os.path.dirname(__file__) + '/files/Sovk_accept.docx')
    doc = replace_doc(doc=doc, words=['{{ФИО}}', '{{Адрес}}', '{{Паспорт}}', '{{Кем_выдан}}',
                                      '{{число}}', '{{месяц}}', '{{год}}'],
                      replacements=[info[3], info[9], info[5], info[7], str(date.day), str(date.month), str(date.year)])
    doc = replace_request(words=['{{ФИО}}', '{число}', '{месяц}', '{год}'],
                          replacements=[info[3], str(date.day), str(date.month), str(date.year)], doc=doc)
    doc.save(os.path.dirname(__file__) + '/files/downloaded.docx')
    return send_from_directory(directory=os.path.abspath(os.path.dirname(__file__) + '/files'),
                               filename='downloaded.docx')

def six(userid):
    date = datetime.today()
    info = read(mode='data')[userid]
    doc = Document(os.path.dirname(__file__) + '/files/Soglasie-TF-2216-191.docx')
    doc = replace_doc(doc=doc, words=['{{ФИО}}', '{{Число_рождения}}', '{{Месяц_рождения}}', '{{Год_рождения}}',
                                      '{{Место_рождения}}', '{{Серия_паспорта}}', '{{Номер}}', '{{Кем_выдан}}',
                                      '{{Адрес}}', '{{Снилс}}', '@', '~', '{{Дата}}'],
                      replacements=[info[3], info[4].split('.')[0], info[4].split('.')[1], info[4].split('.')[2],
                                    info[61], info[5].split(' ')[0], info[5].split(' ')[1], info[7], info[9],
                                    info[8], 'V', '', str(date.day) + '.' + str(date.month) + '.' + str(date.year)])
    doc.save(os.path.dirname(__file__) + '/files/downloaded.docx')
    return send_from_directory(directory=os.path.abspath(os.path.dirname(__file__) + '/files'),
                                   filename='downloaded.docx')

def seven(userid):
    date = datetime.today()
    info = read(mode='data')[userid]
    doc = Document(os.path.dirname(__file__) + '/files/spravka.docx')
    doc = replace_doc(doc=doc, words=['{{число}}', '{{месяц}}', '{{год}}', '{{ФИО}}', '{{ИНН}}', '{{число работы}}', '{{м.работы}}',
                                      '{{г.работы}}', '{{должность}}'],
                      replacements=[str(date.day), str(['января', 'февраля', 'марта', 'апреля', 'мая', 'июня', 'июля', 'августа',
                      'сентября', 'октября', 'ноября', 'декабря'][date.month - 1]), str(date.year)[2:], info[3], info[14],
                                    info[71].split('.')[0], info[71].split('.')[1], info[71].split('.')[2], info[25]])
    years = spravkaFieldPicker(date)
    doc = replace_request(words=['{{наимен орги}}', '{{ИНН}}', '{{реквизиты банка}}', '{{адрес}}', '{{сайт}}', '{{email}}',
                                 '{{телефон}}', '{{м1}}', '{{м2}}', '{{м3}}', '{{м4}}', '{{м5}}', '{{м6}}', '{{г1}}',
                                 '{{г2}}', '{{г3}}', '{{г4}}', '{{г5}}', '{{г6}}', '{{сумма денег}}'],
                          replacements=[info[13], info[14], info[65], info[30], info[33], '', info[31],
                                        years[1][0], years[1][1], years[1][2], years[1][3], years[1][4],
                                        years[1][5], years[0][0], years[0][1], years[0][2], years[0][3],
                                        years[0][4], years[0][5], info[26]], doc=doc)
    doc.save(os.path.dirname(__file__) + '/files/downloaded.docx')
    return send_from_directory(directory=os.path.abspath(os.path.dirname(__file__) + '/files'),
                                   filename='downloaded.docx')

def create(id, userid):
    if int(id) == 4:
        return four(userid)
    if int(id) == 5:
        return five(userid)
    if int(id) == 6:
        return six(userid)
    if int(id) == 7:
        return seven(userid)
